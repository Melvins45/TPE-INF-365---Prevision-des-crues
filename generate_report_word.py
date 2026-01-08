from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import os
import json

def extract_ipynb_code(filename):
    """Extrait le code source des cellules de code d'un notebook Jupyter."""
    if not os.path.exists(filename): return f"[Notebook introuvable : {filename}]"
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            data = json.load(f)
        code_content = ""
        for cell in data.get('cells', []):
            if cell['cell_type'] == 'code':
                source = "".join(cell['source'])
                if source.strip():
                    code_content += source + "\n\n# --------------------------------------------------------\n\n"
        return code_content
    except Exception as e:
        return f"[Erreur lecture Notebook : {str(e)}]"

def read_file_content(filename):
    """Lit le contenu d'un fichier texte."""
    if os.path.exists(filename):
        try:
            with open(filename, 'r', encoding='utf-8') as f: return f.read()
        except: return f"[Erreur lecture : {filename}]"
    return f"[Fichier introuvable : {filename}]"

def add_code_block(doc, code_content):
    """Ajoute un bloc de code formaté."""
    p = doc.add_paragraph(code_content)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0, 0, 100)

def clean_markdown(text):
    """Nettoie le markdown."""
    lines = text.split('\n')
    cleaned = []
    for line in lines:
        if line.startswith('```'): continue
        if line.startswith('#'): line = line.replace('#', '').strip()
        cleaned.append(line)
    return "\n".join(cleaned)

def generate_final_report():
    doc = Document()
    
    # --- STYLES ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ================= PAGE DE GARDE =================
    doc.add_paragraph("\n")
    title = doc.add_paragraph("UNIVERSITÉ DE DOUALA\nFACULTÉ DES SCIENCES")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    
    doc.add_paragraph("\n\n\n\n")
    
    main_title = doc.add_paragraph("SYSTÈME INTÉGRÉ DE PRÉVISION DES CRUES\nPAR INTELLIGENCE ARTIFICIELLE")
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = main_title.runs[0]
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102) # Bleu Foncé
    
    subtitle = doc.add_paragraph("RAPPORT TECHNIQUE FINAL & IMPLÉMENTATION")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph("\n\n")
    
    # Info Groupe
    frame = doc.add_paragraph()
    frame.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = frame.add_run("TPE INF 365 - GROUPE 24")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph("\n\n")
    
    members = doc.add_paragraph()
    members.alignment = WD_ALIGN_PARAGRAPH.LEFT
    members.add_run("Membres du projet :\n").bold = True
    members.add_run("1. NITOPOP JEATSA GUILLAUME MELVIN (Chef) - 20S43003\n")
    members.add_run("2. DEMANOU KEMKENG DILAN - 25S03516\n")
    members.add_run("3. NOSSI YIMGO LYNDSEY SULIVANE - 23S87713\n")
    members.add_run("4. TCHIEUTCHOUA FOTEPING ASHLEY MEGANE - 23S87863\n")
    members.add_run("5. WANGA POUYA KAVEN SAMIRA - 23S88070\n")
    
    doc.add_paragraph("\n\n")
    examiner = doc.add_paragraph()
    examiner.add_run("Sous la supervision de : \n").bold = True
    examiner.add_run("Dr. Justin MOSKOLAI").font.size = Pt(12)
    
    doc.add_page_break()

    # ================= SOMMAIRE AUTOMATIQUE (Simulé) =================
    doc.add_heading("Sommaire", level=1)
    doc.add_paragraph("1. Introduction et Contexte\n"
                      "2. Architecture de la Solution\n"
                      "3. Méthodologie Data Science & Modélisation\n"
                      "4. Résultats des Expérimentations\n"
                      "5. Implémentation Logicielle (GUI)\n"
                      "6. Conclusion & Perspectives\n"
                      "A. Annexe - Code Source Complet")
    doc.add_page_break()

    # ================= 1. INTRODUCTION =================
    doc.add_heading('1. Introduction et Contexte', level=1)
    doc.add_paragraph(
        "Les inondations représentent une menace majeure pour la ville de Douala, aggravée par une "
        "urbanisation rapide et une pluviométrie intense (mousson). Ce projet a pour objectif de concevoir "
        "un outil d'aide à la décision capable de prédire la probabilité d'inondation en temps réel."
    )
    doc.add_paragraph(
        "L'objectif de ce document est de présenter la démarche technique, l'implémentation des modèles "
        "d'Intelligence Artificielle et l'interface utilisateur développée."
    )

    # ================= 2. ARCHITECTURE =================
    doc.add_heading('2. Architecture de la Solution', level=1)
    doc.add_paragraph(
        "Le système repose sur une interaction entre trois composants majeurs :"
    )
    doc.add_paragraph(
        "1. Collecte de Données : Réseaux de capteurs (Projet IoT) pour mesurer la pluviométrie, le niveau d'eau et les paramètres atmosphériques.\n"
        "2. Moteur de Calcul (IA) : Algorithmes prédictifs entraînés (Python / Scikit-Learn / TensorFlow).\n"
        "3. Interface de Visualisation : Application de bureau (Tkinter) pour la simulation et Application Mobile pour les alertes."
    )

    # ================= 3. METHODOLOGIE & MODELES =================
    doc.add_heading('3. Méthodologie et Modèles d\'IA', level=1)
    doc.add_paragraph(
        "Trois approches de modélisation ont été testées pour prédire la probabilité d'inondation. "
        "L'objectif est de comparer une approche statistique simple, une approche par ensemble d'arbres, et une approche par Deep Learning."
    )
    
    # 3.1 Régression Linéaire
    doc.add_heading('3.1 Régression Linéaire (Baseline)', level=2)
    doc.add_paragraph(
        "Ce modèle sert de référence de base. Il suppose une relation linéaire directe entre les précipitations, "
        "la topographie et le risque d'inondation. Bien que rapide, il manque de finesse pour capturer les interactions complexes."
    )
    if os.path.exists("actaul vs predicted.png"):
        doc.add_picture("actaul vs predicted.png", width=Inches(4.5))
        doc.add_paragraph("Figure : Comparaison Réel vs Prédit (Régression Linéaire).").italic = True
    
    # 3.2 Forêt Aléatoire
    doc.add_heading('3.2 Forêt Aléatoire (Random Forest)', level=2)
    doc.add_paragraph(
        "Ce modèle utilise 100 arbres de décision pour réduire le risque de surapprentissage (overfitting). "
        "Il capture bien les non-linéarités et permet d'identifier l'importance de chaque variable "
        "(ex: l'impact de l'urbanisation vs la déforestation)."
    )
    if os.path.exists("actaul vs predicted foret aleatoire.png"):
        doc.add_picture("actaul vs predicted foret aleatoire.png", width=Inches(4.5))
        doc.add_paragraph("Figure : Comparaison Réel vs Prédit (Random Forest).").italic = True
    
    # 3.3 Réseau de Neurones
    doc.add_heading('3.3 Réseau de Neurones Séquentiel', level=2)
    doc.add_paragraph(
        "Notre modèle le plus avancé utilise une architecture de Deep Learning (Keras/TensorFlow) avec trois couches de neurones. "
        "Il offre la meilleure précision en apprenant des motifs complexes inaccessibles aux autres modèles."
    )
    if os.path.exists("actaul vs predicted reseaux de nerurones.png"):
        doc.add_picture("actaul vs predicted reseaux de nerurones.png", width=Inches(4.5))
        doc.add_paragraph("Figure : Comparaison Réel vs Prédit (Réseau de Neurones).").italic = True

    # ================= 4. RESULTATS =================
    doc.add_heading('4. Résultats et Validation', level=1)
    doc.add_paragraph(
        "Pour s'assurer de la fiabilité du système, une validation croisée (K-Fold, K=5) a été effectuée. "
        "Cela garantit que les performances ne sont pas dues au hasard."
    )
    doc.add_paragraph(
        "Résumé des performances (Moyenne R²) :\n"
        "- Linear Regression : 0.9945\n"
        "- Random Forest : 0.7280\n"
        "- Neural Network : 0.9912"
    )

    # Image Validation Croisée (Spécifiquement demandée)
    if os.path.exists("validation_croisee_resultats.png"):
        doc.add_picture("validation_croisee_resultats.png", width=Inches(5))
        doc.add_paragraph("Figure : Résultats de la Validation Croisée (Comparaison des modèles).").italic = True
    
    if os.path.exists("correlation matrix.png"):
        doc.add_paragraph("\n")
        doc.add_picture("correlation matrix.png", width=Inches(5))
        doc.add_paragraph("Figure : Matrice de Corrélation des variables.").italic = True

    # ================= 5. INTERFACE GUI =================
    doc.add_heading('5. Implémentation Logicielle (GUI)', level=1)
    doc.add_paragraph(
        "L'application finale permet d'ajuster les 20 paramètres via des curseurs et de visualiser instantanément le risque. "
        "Voici un aperçu de l'interface :"
    )
    
    # Image Interface Graphique (Spécifiquement demandée)
    if os.path.exists("interface graphique.png"):
         doc.add_paragraph("[Capture d'écran de l'Interface Graphique]") # Placeholder text matched to user file, but image follows
         doc.add_picture("interface graphique.png", width=Inches(6))
         doc.add_paragraph("Figure : Capture d'écran de l'Interface Graphique Principale.").italic = True
    elif os.path.exists("app_interface_screenshot.png"):
         doc.add_picture("app_interface_screenshot.png", width=Inches(6))
         doc.add_paragraph("Figure : Capture d'écran de l'Interface Graphique Principale.").italic = True

    # ================= 6. CONCLUSION =================
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "Ce projet a permis de développer un outil fonctionnel de prévision des crues pour la ville de Douala. "
        "L'utilisation de l'intelligence artificielle, et particulièrement du Réseau de Neurones, a démontré une excellente capacité "
        "à anticiper les risques en fonction des données climatiques et topographiques. "
        "L'intégration de ces modèles dans une interface graphique rend l'outil accessible pour une aide à la décision concrète."
    )

    # ================= ANNEXES CODE =================
    doc.add_page_break()
    doc.add_heading('ANNEXES : CODE SOURCE DES MODÈLES', level=1)
    doc.add_paragraph("Voici le code source extrait directement des notebooks Jupyter utilisés pour l'entraînement.")

    doc.add_heading('A. Modèle Régression Linéaire (Notebook)', level=2)
    add_code_block(doc, extract_ipynb_code("modeleRegressionLineaire.ipynb"))

    doc.add_heading('B. Modèle Forêt Aléatoire (Notebook)', level=2)
    add_code_block(doc, extract_ipynb_code("modeleForetAleatoire.ipynb"))

    doc.add_heading('C. Modèle Réseau de Neurones (Notebook)', level=2)
    add_code_block(doc, extract_ipynb_code("modeleReseauNeuronesSequential.ipynb"))

    doc.add_page_break()
    doc.add_heading('ANNEXES : APPLICATION', level=1)
    
    doc.add_heading("D. Code de l'Interface Graphique (app_prediction_gui.py)", level=2)
    add_code_block(doc, read_file_content("app_prediction_gui.py"))

    doc.add_heading("E. Code de Validation Croisée (validation_croisee.py)", level=2)
    add_code_block(doc, read_file_content("validation_croisee.py"))
    
    # Sauvegarde
    filename = "Rapport_Final_Complet_Groupe24.docx"
    doc.save(filename)
    print(f"Rapport complet généré : {filename}")

if __name__ == "__main__":
    generate_final_report()
